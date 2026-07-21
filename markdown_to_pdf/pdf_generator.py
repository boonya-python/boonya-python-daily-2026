import os
import time
import uuid
import shutil


class PdfGenerator:
    def generate(self, html_content, output_path):
        generators = [
            (self._from_pdfkit, html_content, output_path),
            (self._from_docx2pdf, html_content, output_path),
            (self._from_win32com, html_content, output_path),
        ]

        result = False
        for gen, src, dst in generators:
            try:
                gen(src, dst)
                if os.path.exists(dst) and os.path.getsize(dst) > 0:
                    print(f'✅ PDF 已生成: {dst}')
                    result = True
                    break
            except Exception as e:
                print(f"⚠️ {gen.__name__} 失败: {e}，尝试下一个方案")

        if not result:
            print("❌ 所有 PDF 生成方案均失败")
        return result

    def _from_pdfkit(self, html_content, pdf_path):
        import pdfkit

        full_html = f'''
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: "Microsoft YaHei", "PingFang SC", Arial, sans-serif; padding: 30px; line-height: 1.8; }}
                img {{ max-width: 100%; height: auto; display: block; margin: 10px 0; }}
                pre, code {{ background: #f4f4f4; padding: 10px; border-radius: 4px; overflow-x: auto; font-family: "Consolas", "Monaco", monospace; font-size: 12px; }}
                h1, h2, h3 {{ margin-top: 20px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        '''

        options = {
            'enable-local-file-access': None,
            'encoding': 'UTF-8',
            'page-size': 'A4',
            'margin-top': '15mm',
            'margin-bottom': '15mm',
            'margin-left': '15mm',
            'margin-right': '15mm',
        }
        pdfkit.from_string(full_html, pdf_path, options=options)

    def _from_docx2pdf(self, html_content, pdf_path):
        from docx import Document
        from docx.shared import Inches, Pt
        from docx2pdf import convert
        from bs4 import BeautifulSoup

        timestamp = time.strftime('%Y%m%d_%H%M%S')
        unique_id = uuid.uuid4().hex[:6]
        temp_docx = pdf_path.replace('.pdf', f'_temp_{timestamp}_{unique_id}.docx')

        doc = Document()

        soup = BeautifulSoup(html_content, 'html.parser')

        def process_element(elem):
            if elem is None:
                return
            if isinstance(elem, str):
                text = elem.strip()
                if text:
                    doc.add_paragraph(text)
                return

            tag = elem.name
            if tag is None:
                return

            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                text = elem.get_text().strip()
                if text:
                    doc.add_heading(text, level=level)
                return

            if tag == 'p':
                text = elem.get_text().strip()
                if text:
                    doc.add_paragraph(text)
                for img in elem.find_all('img'):
                    process_element(img)
                return

            if tag == 'pre':
                code = elem.get_text().strip()
                if code:
                    p = doc.add_paragraph()
                    run = p.add_run(code)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                return

            if tag == 'img':
                src = elem.get('src')
                if src and os.path.exists(src):
                    try:
                        doc.add_picture(src, width=Inches(6))
                    except Exception as e:
                        print(f'  ⚠️ 插入图片失败: {src} - {e}')
                return

            if tag in ['ul', 'ol']:
                for li in elem.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet' if tag == 'ul' else 'List Number')
                return

            if tag == 'table':
                rows = elem.find_all('tr')
                if rows:
                    first_row = rows[0]
                    cols = len(first_row.find_all(['th', 'td']))

                    table = doc.add_table(rows=1, cols=cols)
                    hdr_cells = table.rows[0].cells
                    for i, cell in enumerate(first_row.find_all(['th', 'td'])):
                        hdr_cells[i].text = cell.get_text().strip()

                    for row in rows[1:]:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(row.find_all(['th', 'td'])):
                            if i < len(row_cells):
                                row_cells[i].text = cell.get_text().strip()
                return

            for child in elem.children:
                process_element(child)

        if soup.body:
            for child in soup.body.children:
                process_element(child)
        else:
            for child in soup.children:
                process_element(child)

        doc.save(temp_docx)

        convert(temp_docx, pdf_path)

        if os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except:
                pass

    def _from_win32com(self, html_content, pdf_path):
        import win32com.client
        from docx import Document
        from docx.shared import Inches, Pt
        from bs4 import BeautifulSoup

        timestamp = time.strftime('%Y%m%d_%H%M%S')
        unique_id = uuid.uuid4().hex[:6]
        temp_docx = pdf_path.replace('.pdf', f'_temp_{timestamp}_{unique_id}.docx')

        doc = Document()

        soup = BeautifulSoup(html_content, 'html.parser')

        def process_element(elem):
            if elem is None:
                return
            if isinstance(elem, str):
                text = elem.strip()
                if text:
                    doc.add_paragraph(text)
                return

            tag = elem.name
            if tag is None:
                return

            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                text = elem.get_text().strip()
                if text:
                    doc.add_heading(text, level=level)
                return

            if tag == 'p':
                text = elem.get_text().strip()
                if text:
                    doc.add_paragraph(text)
                for img in elem.find_all('img'):
                    process_element(img)
                return

            if tag == 'pre':
                code = elem.get_text().strip()
                if code:
                    p = doc.add_paragraph()
                    run = p.add_run(code)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                return

            if tag == 'img':
                src = elem.get('src')
                if src and os.path.exists(src):
                    try:
                        doc.add_picture(src, width=Inches(6))
                    except Exception as e:
                        print(f'  ⚠️ 插入图片失败: {src} - {e}')
                return

            if tag in ['ul', 'ol']:
                for li in elem.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet' if tag == 'ul' else 'List Number')
                return

            if tag == 'table':
                rows = elem.find_all('tr')
                if rows:
                    first_row = rows[0]
                    cols = len(first_row.find_all(['th', 'td']))

                    table = doc.add_table(rows=1, cols=cols)
                    hdr_cells = table.rows[0].cells
                    for i, cell in enumerate(first_row.find_all(['th', 'td'])):
                        hdr_cells[i].text = cell.get_text().strip()

                    for row in rows[1:]:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(row.find_all(['th', 'td'])):
                            if i < len(row_cells):
                                row_cells[i].text = cell.get_text().strip()
                return

            for child in elem.children:
                process_element(child)

        if soup.body:
            for child in soup.body.children:
                process_element(child)
        else:
            for child in soup.children:
                process_element(child)

        doc.save(temp_docx)

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(temp_docx))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
        finally:
            word.Quit()

        if os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except:
                pass