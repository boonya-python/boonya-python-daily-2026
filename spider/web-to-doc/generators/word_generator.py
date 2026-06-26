import os
import re
from docx import Document
from docx.shared import Inches, Pt

from datetime import datetime


def _is_css_style_code(text):
    css_keywords = ['!important', 'rgba(', '{', '}', ';', 'px;', 'em;', '%',
                    'color:', 'font-', 'background:', 'margin:', 'padding:',
                    'border:', 'display:', 'position:', 'width:', 'height:']
    count = sum(1 for kw in css_keywords if kw in text)
    if count >= 5 and text.count('{') > 2 and text.count('}') > 2:
        return True
    if re.match(r'^\.[\w-]+\s*\{', text.strip()):
        return True
    if re.match(r'^pre\s*\{|^code\s*\{|^table\s*\{', text.strip()):
        return True
    return False


class WordGenerator:
    def generate(self, title, content_div, url_to_local, output_path, source_url=None):
        doc = Document()

        # 添加来源信息头部
        if source_url:
            doc.add_paragraph('─' * 50)
            doc.add_paragraph(f'📌 来源: {source_url}')
            doc.add_paragraph(f'📅 采集时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('─' * 50)
            doc.add_paragraph()

        doc.add_heading(title, level=1)

        def process_element(elem):
            if elem is None:
                return
            if isinstance(elem, str):
                text = elem.strip()
                if text:
                    doc.add_paragraph(text)
                return

            tag = elem.name
            if tag is None:
                return

            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                text = elem.get_text().strip()
                if text:
                    doc.add_heading(text, level=level)
                return

            if tag == 'p':
                img = elem.find('img')
                if img and len(elem.get_text(strip=True)) == 0:
                    process_element(img)
                    return
                text = elem.get_text().strip()
                if text:
                    doc.add_paragraph(text)
                return

            if tag == 'pre':
                code = elem.get_text().strip()
                if code and not _is_css_style_code(code):
                    p = doc.add_paragraph()
                    run = p.add_run(code)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                return

            if tag == 'img':
                src = elem.get('src') or elem.get('data-src')
                if src:
                    local_path = url_to_local.get(src)
                    if local_path and os.path.exists(local_path):
                        try:
                            doc.add_picture(local_path, width=Inches(6))
                        except Exception as e:
                            print(f'  ⚠️ 插入图片失败: {src} - {e}')
                    else:
                        print(f'  ⚠️ 未找到图片本地路径: {src[:60]}')
                return

            if tag in ['ul', 'ol']:
                for li in elem.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(
                            text,
                            style='List Bullet' if tag == 'ul' else 'List Number'
                        )
                return

            if tag == 'table':
                text = elem.get_text().strip()
                if text:
                    doc.add_paragraph(text)
                return

            for child in elem.children:
                process_element(child)

        if content_div is None:
            print("❌ 未找到文章正文，Word 文档仅包含标题")
        else:
            process_element(content_div)

        doc.save(output_path)
        print(f'✅ Word 已生成: {output_path}')
