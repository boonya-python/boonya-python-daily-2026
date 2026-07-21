#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
CSDN 文章抓取工具 —— 导出为 Word 和 PDF（含图片）
使用 Selenium 渲染动态内容，支持懒加载图片。
"""

import os
import time
import shutil
from copy import deepcopy
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from docx.shared import Inches, Pt
import pdfkit
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager

# ========== 配置 ==========
HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}
IMG_DIR = './temp_images'          # 临时图片存放目录
os.makedirs(IMG_DIR, exist_ok=True)


def fetch_html_with_selenium(url):
    """
    使用 Selenium + Chrome 获取渲染后的完整 HTML
    已添加忽略 SSL 证书错误的参数，解决 ERR_SSL_PROTOCOL_ERROR
    """
    chrome_options = Options()
    chrome_options.add_argument('--headless')               # 无头模式
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument(f'user-agent={HEADERS["User-Agent"]}')

    # ===== SSL 错误修复 =====
    chrome_options.add_argument('--ignore-certificate-errors')
    chrome_options.add_argument('--ignore-ssl-errors')
    chrome_options.add_argument('--allow-insecure-localhost')
    chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])

    driver = webdriver.Chrome(
        service=Service(ChromeDriverManager().install()),
        options=chrome_options
    )

    try:
        driver.get(url)
        # 等待页面核心元素加载（可调整）
        time.sleep(3)
        # 滚动至底部，触发图片懒加载
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(1.5)                      # 等待图片加载
        html = driver.page_source
    except Exception as e:
        print(f"❌ Selenium 获取页面失败: {e}")
        html = None
    finally:
        driver.quit()

    return html


def parse_article(html, base_url):
    """
    从 HTML 中提取文章标题、正文容器和所有图片 URL
    支持 CSDN 多种常见类名/ID
    """
    soup = BeautifulSoup(html, 'lxml')

    # ---------- 提取标题 ----------
    title_tag = soup.find('h1', class_='title-article')
    if not title_tag:
        title_tag = soup.find('h1', class_='blog-title')
    if not title_tag:
        title_tag = soup.find('h1')
    title = title_tag.get_text().strip() if title_tag else '未命名文章'

    # ---------- 提取正文容器 ----------
    content_div = soup.find('div', id='article_content')
    if not content_div:
        content_div = soup.find('div', class_='article_content')
    if not content_div:
        content_div = soup.find('div', id='content_views')
    if not content_div:
        # 兜底：找包含最多文本的 div
        all_divs = soup.find_all('div')
        if all_divs:
            content_div = max(all_divs, key=lambda d: len(d.get_text(strip=True)))

    # ---------- 提取图片 URL ----------
    img_urls = []
    if content_div:
        img_tags = content_div.find_all('img')
        for img in img_tags:
            src = img.get('src') or img.get('data-src')
            if src and not src.startswith('data:'):   # 忽略 base64 图片
                full_url = urljoin(base_url, src)
                img_urls.append(full_url)

    return title, content_div, img_urls


def download_image(img_url, save_path):
    """下载单张图片，返回是否成功"""
    try:
        resp = requests.get(img_url, headers=HEADERS, timeout=10)
        if resp.status_code == 200:
            with open(save_path, 'wb') as f:
                f.write(resp.content)
            return True
    except Exception as e:
        print(f'  ⚠️ 下载失败: {img_url[:60]}... {e}')
    return False


def download_all_images(img_urls):
    """
    批量下载图片，返回：
        local_paths : 本地路径列表
        url_to_local: dict {原始URL: 本地路径}
    """
    local_paths = []
    url_to_local = {}
    for i, url in enumerate(img_urls):
        # 提取扩展名
        ext = os.path.splitext(url.split('?')[0])[1]
        if not ext or ext.lower() not in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp']:
            ext = '.jpg'
        filename = f'img_{i+1:03d}{ext}'
        local_path = os.path.join(IMG_DIR, filename)
        print(f'  📥 下载图片 {i+1}/{len(img_urls)}: {filename}')
        if download_image(url, local_path):
            local_paths.append(local_path)
            url_to_local[url] = local_path
        time.sleep(0.3)   # 避免请求过快
    return local_paths, url_to_local


def generate_word(title, content_div, url_to_local, output_path):
    """
    递归遍历 content_div，按顺序生成 Word 文档
    支持标题、段落、代码块、列表、表格和图片
    """
    doc = Document()
    doc.add_heading(title, level=1)

    def process_element(elem):
        if elem is None:
            return
        if isinstance(elem, str):
            text = elem.strip()
            if text:
                doc.add_paragraph(text)
            return

        tag = elem.name
        if tag is None:
            return

        # ----- 标题 -----
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            text = elem.get_text().strip()
            if text:
                doc.add_heading(text, level=level)
            return

        # ----- 段落（含内部图片） -----
        if tag == 'p':
            img = elem.find('img')
            if img and len(elem.get_text(strip=True)) == 0:
                process_element(img)
                return
            text = elem.get_text().strip()
            if text:
                doc.add_paragraph(text)
            return

        # ----- 代码块 -----
        if tag == 'pre':
            code = elem.get_text().strip()
            if code:
                p = doc.add_paragraph()
                run = p.add_run(code)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            return

        # ----- 图片 -----
        if tag == 'img':
            src = elem.get('src') or elem.get('data-src')
            if src:
                local_path = url_to_local.get(src)
                if local_path and os.path.exists(local_path):
                    try:
                        doc.add_picture(local_path, width=Inches(6))
                    except Exception as e:
                        print(f'  ⚠️ 插入图片失败: {src} - {e}')
                else:
                    print(f'  ⚠️ 未找到图片本地路径: {src[:60]}')
            return

        # ----- 列表 -----
        if tag in ['ul', 'ol']:
            for li in elem.find_all('li', recursive=False):
                text = li.get_text().strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet' if tag == 'ul' else 'List Number')
            return

        # ----- 表格（简单提取文本） -----
        if tag == 'table':
            text = elem.get_text().strip()
            if text:
                doc.add_paragraph(text)
            return

        # ----- 其他容器，递归处理子元素 -----
        for child in elem.children:
            process_element(child)

    if content_div is None:
        print("❌ 未找到文章正文，Word 文档仅包含标题")
    else:
        process_element(content_div)

    doc.save(output_path)
    print(f'✅ Word 已生成: {output_path}')


def generate_pdf(content_div, url_to_local, output_path):
    """
    使用 pdfkit 生成 PDF（需 wkhtmltopdf）
    先将图片 src 替换为本地路径，再生成
    """
    if content_div is None:
        print("❌ 无法生成 PDF：正文为空")
        return

    # 复制 content_div，修改 img 的 src 为本地路径
    div_copy = deepcopy(content_div)
    for img in div_copy.find_all('img'):
        src = img.get('src') or img.get('data-src')
        if src and src in url_to_local:
            img['src'] = url_to_local[src]
        elif src:
            # 若无法映射，保留原 URL（可能无法加载）
            pass

    # 构建完整 HTML
    full_html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: "Microsoft YaHei", "PingFang SC", Arial, sans-serif; padding: 30px; line-height: 1.8; }}
            img {{ max-width: 100%; height: auto; display: block; margin: 10px 0; }}
            pre, code {{ background: #f4f4f4; padding: 10px; border-radius: 4px; overflow-x: auto; }}
            h1, h2, h3 {{ margin-top: 20px; }}
        </style>
    </head>
    <body>
        {str(div_copy)}
    </body>
    </html>
    '''

    try:
        options = {
            'enable-local-file-access': None,
            'encoding': 'UTF-8',
            'page-size': 'A4',
            'margin-top': '15mm',
            'margin-bottom': '15mm',
            'margin-left': '15mm',
            'margin-right': '15mm',
        }
        pdfkit.from_string(full_html, output_path, options=options)
        print(f'✅ PDF 已生成: {output_path}')
    except Exception as e:
        print(f'❌ PDF 生成失败: {e}')
        print('   请确认 wkhtmltopdf 已正确安装并添加到 PATH')


def main():
    # 可替换为任意 CSDN 文章 URL
    url = 'https://blog.csdn.net/weixin_44566643/article/details/107552945'
    print(f'🌐 正在抓取（使用 Selenium）: {url}')

    # 1. 获取渲染后的 HTML
    html = fetch_html_with_selenium(url)
    if not html:
        print("❌ 无法获取页面，程序退出")
        return

    # 2. 解析文章
    title, content_div, img_urls = parse_article(html, url)
    print(f'📄 标题: {title}')
    print(f'🖼️  发现 {len(img_urls)} 张图片')

    # 3. 下载图片
    img_paths, url_to_local = download_all_images(img_urls)
    print(f'✅ 成功下载 {len(img_paths)} 张图片')

    # 4. 生成 Word
    # 清理文件名中的非法字符（可选）
    safe_title = "".join(c for c in title if c not in r'\/:*?"<>|')
    word_path = f'./{safe_title}.docx'
    generate_word(title, content_div, url_to_local, word_path)

    # 5. 生成 PDF（若无需 PDF，可注释此行）
    pdf_path = f'./{safe_title}.pdf'
    generate_pdf(content_div, url_to_local, pdf_path)

    # 6. 清理临时图片（可选）
    # shutil.rmtree(IMG_DIR, ignore_errors=True)

    print('\n🎉 完成！')


if __name__ == '__main__':
    main()